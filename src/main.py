import docx
from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
import logging
import io
import os
import sys
import tempfile
import configparser
from PIL import Image as PILImage
import argparse

from utils.config import ConfigManager
from docx import Document
from document.formatter import DocumentFormatter  # 添加 DocumentFormatter 的导入

# Configure logging to display messages with timestamps and severity levels
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('file_format.log', encoding='utf-8')
    ]
)

def load_config(config_file):
    """Load configuration settings from the specified INI file."""
    config_manager = ConfigManager(config_file)
    return config_manager.load()

def format_word(input_file, output_file, config):
    """Format the Word document according to the specified configuration."""
    # 验证输入文件
    if not os.path.exists(input_file):
        raise FileNotFoundError(f"输入文件不存在: {input_file}")
    
    # 验证文件扩展名
    if not input_file.lower().endswith('.docx'):
        raise ValueError(f"不支持的文件格式，仅支持.docx文件: {input_file}")
    
    # 验证输出目录是否存在
    output_dir = os.path.dirname(output_file)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    try:
        doc = Document(input_file)
        new_doc = Document()
        formatter = DocumentFormatter(config)
        
        # 处理表格
        for table in doc.tables:
            formatter.copy_table(table, new_doc)
        
        # 处理段落
        order_number = 1
        for index, paragraph in enumerate(doc.paragraphs):
            logging.info(f'处理段落 {index + 1}/{len(doc.paragraphs)}: {paragraph.text}')
            
            if paragraph.style.name.startswith('Heading'):
                level = int(paragraph.style.name[-1])
                logging.info(f'发现标题级别 {level}: {paragraph.text}')
                formatter.format_heading(paragraph, new_doc, level)
                order_number = 1
            
            elif paragraph.style.name == 'Normal':
                logging.info(f'格式化普通段落: {paragraph.text}')
                _, order_number = formatter.format_normal_paragraph(paragraph, new_doc, order_number)
        
        # 保存文档
        logging.info(f'保存格式化文档: {output_file}')
        new_doc.save(output_file)
        
    except Exception as e:
        logging.error(f"处理文档失败: {str(e)}")
        raise

def main():
    parser = argparse.ArgumentParser(description='处理Word文档格式化')
    parser.add_argument('--version', action='version', version='FileFormat 1.0.0')
    parser.add_argument('--input', required=True, help='输入Word文档路径')
    parser.add_argument('--output', required=True, help='输出Word文档路径')
    parser.add_argument('--config', default='config.ini', help='配置文件路径')
    
    args = parser.parse_args()
    try:
        config = load_config(args.config)
        format_word(args.input, args.output, config)
        logging.info("文档处理完成")
    except Exception as e:
        logging.error(f"处理失败: {str(e)}")
        sys.exit(1)

if __name__ == '__main__':
    main()