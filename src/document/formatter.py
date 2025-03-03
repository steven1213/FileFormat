from docx import Document
from docx.shared import Pt
import logging
import io
from .image import process_image
from docx.shared import Inches
from docx.oxml import parse_xml  # 添加这行导入
from .table_handler import TableHandler

class DocumentFormatter:
    def __init__(self, config):
        self.config = config
        self.table_handler = TableHandler()
        self._init_typography_settings()

    def _init_typography_settings(self):
        self.body_font_name = self.config['Typography']['body_font'].strip('"')
        self.body_font_size = Pt(int(self.config['Typography']['body_size']))
        self.first_line_indent = Pt(int(self.config['Typography']['first_line_indent']))
        self.line_spacing = float(self.config['Typography']['line_spacing'])
        self.min_heading_level = int(self.config['Typography']['min_heading_level'])

    def format_heading(self, paragraph, new_doc, level):
        try:
            title_font_name = self.config['Typography'][f'title_level_{level}_font'].strip('"')
            title_font_size = Pt(int(self.config['Typography'][f'title_level_{level}_size']))
        except KeyError:
            logging.warning(f'No configuration found for title level {level}. Using default font and size.')
            title_font_name = self.body_font_name
            title_font_size = self.body_font_size

        new_paragraph = new_doc.add_paragraph()
        new_paragraph.style = f'Heading {level}'  # 设置正确的标题样式
        
        for run in paragraph.runs:
            new_run = new_paragraph.add_run(run.text)
            new_run.font.name = title_font_name
            new_run.font.size = title_font_size
            new_run.font.bold = True  # 标题加粗
        
        return new_paragraph

    def format_normal_paragraph(self, paragraph, new_doc, order_number):
        new_paragraph = new_doc.add_paragraph()
        new_paragraph.style = 'Normal'
        
        # 设置段落格式
        paragraph_format = new_paragraph.paragraph_format
        paragraph_format.first_line_indent = self.first_line_indent
        paragraph_format.line_spacing = self.line_spacing  # 直接使用浮点数
        paragraph_format.space_after = Pt(0)  # 段后间距设为0
        
        # 复制内容并设置格式
        for run in paragraph.runs:
            new_run = new_paragraph.add_run(run.text)
            new_run.font.name = self.body_font_name
            new_run.font.size = self.body_font_size
            
            if run._element.xpath('.//a:blip'):
                process_image(run, new_paragraph, self.config)
        
        # 检查是否需要编号
        if self._should_add_numbering(new_paragraph.text):
            return new_paragraph, self._apply_numbering(new_paragraph, order_number)
        return new_paragraph, order_number

    def _should_add_numbering(self, text):
        """检查段落是否需要编号"""
        # 定义不需要编号的标识符或条件
        no_number_markers = [
            '不需要编号的标识',
            '注：',
            '说明：',
            '备注：',
            # 添加其他不需要编号的标识
        ]
        return not any(text.strip().startswith(marker) for marker in no_number_markers)
    def _apply_numbering(self, paragraph, order_number):
        if paragraph.text.strip() and not paragraph.text.startswith('不需要编号的标识'):
            if order_number <= 5:
                if order_number <= 9:
                    paragraph.text = f'({order_number}) {paragraph.text}'
                elif order_number == 10:
                    paragraph.text = f'{order_number}) {paragraph.text}'
                else:
                    paragraph.text = f'{chr(96 + (order_number - 1))}) {paragraph.text}'
                return order_number + 1
        return order_number

    def copy_table(self, source_table, target_doc):
        # 创建新表格并复制基本属性
        new_table = target_doc.add_table(rows=len(source_table.rows), cols=len(source_table.columns))
        new_table.style = source_table.style
        new_table.alignment = source_table.alignment
        
        # 复制每个单元格的内容和样式
        for row_idx, row in enumerate(source_table.rows):
            new_row = new_table.rows[row_idx]
            self.table_handler.copy_row_height(row, new_row)
            
            for cell_idx, cell in enumerate(row.cells):
                new_cell = new_row.cells[cell_idx]
                self.table_handler.copy_cell_properties(cell, new_cell)
                self._copy_cell_content(cell, new_cell)
        
        return new_table

    def _copy_cell_content(self, source_cell, target_cell):
        for paragraph in source_cell.paragraphs:
            new_paragraph = target_cell.add_paragraph()
            if paragraph._p.pPr is not None:
                new_paragraph._p.get_or_add_pPr().append(paragraph._p.pPr)
            
            for run in paragraph.runs:
                new_run = new_paragraph.add_run(run.text)
                self._copy_run_format(run, new_run)

    def _copy_run_format(self, source_run, target_run):
        if source_run.font.name:
            target_run.font.name = source_run.font.name
        if source_run.font.size:
            target_run.font.size = source_run.font.size
        target_run.font.bold = source_run.font.bold
        target_run.font.italic = source_run.font.italic
        target_run.font.underline = source_run.font.underline
        
        if source_run._element.xpath('.//a:blip'):
            self.table_handler.copy_image(source_run, target_run)

    def _copy_image_in_cell(self, source_run, target_run):
        """复制单元格中的图片并保持原始格式"""
        for blip in source_run._element.xpath('.//a:blip'):
            rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if rId:
                image_part = source_run.part.rels[rId].target_part
                image_stream = io.BytesIO(image_part.blob)
                
                # 获取原始图片属性
                inline = source_run._element.xpath('.//wp:inline')[0]
                width = int(inline.extent.cx)
                height = int(inline.extent.cy)
                
                # 使用原始尺寸添加图片
                target_run.add_picture(
                    image_stream,
                    width=width,
                    height=height
                )